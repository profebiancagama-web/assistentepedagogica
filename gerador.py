import streamlit as st
from docx import Documentimport streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from google import genai
from google.genai import types
from pypdf import PdfReader
from supabase import create_client, Client
import io, os, re

st.set_page_config(page_title="Gerador Olegário Pro", layout="wide")

SUPABASE_URL = "https://hgdpahtyvoridomkktkt.supabase.co"
SUPABASE_KEY = "sb_publishable_yxUHijFIrK7h2yznh8ySKw_4eBFxAlI"
MINHA_CHAVE = "AQ.Ab8RN6LW-l2a1Qvu-8ZCD75G0OlT5RSVhxJpS3Pa6VcNQ88AFg"

if "supabase" not in st.session_state: st.session_state.supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
if "autenticado" not in st.session_state: st.session_state.autenticado = False
if "client" not in st.session_state: st.session_state.client = genai.Client(api_key=MINHA_CHAVE)
if "resultado" not in st.session_state: st.session_state.resultado = ""
if "historico_chat" not in st.session_state:
    st.session_state.historico_chat = [
        {"role": "model", "text": "Olá! Eu sou a B.IA, sua assistente pedagógica e companheira de trabalho. Sobre o que você gostaria de conversar ou planejar hoje? Estou pronta para ajudar em qualquer assunto!"}
    ]

# --- TELA DE LOGIN ---
if not st.session_state.autenticado:
    st.title("🔐 B.IA: Assistente Pedagógica")
    aba_login, aba_cadastro = st.tabs(["Entrar no Sistema", "Cadastrar Novo Professor"])
    with aba_login:
        email = st.text_input("E-mail funcional:", key="login_email")
        senha = st.text_input("Senha de acesso:", type="password", key="login_senha")
        if st.button("🔓 VALIDAR CREDENCIAIS", key="btn_login_exec"):
            try:
                res = st.session_state.supabase.auth.sign_in_with_password({"email": email, "password": senha})
                if res.session:
                    st.session_state.autenticado = True
                    st.success("✅ Credenciais validadas com sucesso!")
            except:
                st.error("E-mail ou senha incorretos no banco de dados.")
        if st.session_state.autenticado:
            if st.button("🚀 CLIQUE AQUI PARA ENTRAR NO PAINEL"):
                st.rerun()
    with aba_cadastro:
        novo_email = st.text_input("E-mail do Professor:", key="cad_email")
        nova_senha = st.text_input("Crie uma Senha (mínimo 6 dígitos):", type="password", key="cad_senha")
        if st.button("📝 CADASTRAR PROFESSOR", key="btn_cad_exec"):
            try:
                st.session_state.supabase.auth.sign_up({"email": novo_email, "password": nova_senha})
                st.success("Professor cadastrado com sucesso! Volte à aba anterior.")
            except Exception as e: st.error(f"Erro ao cadastrar: {e}")
    st.stop()

# --- INTERFACE PRINCIPAL ---
st.title("🧠 B.IA: Assistente Pedagógica")
with st.sidebar:
    st.header("⚙️ Painel do Professor")
    net = st.toggle("🔓 Liberar Internet (Google Search)", value=False)
    st.write("---")
    if st.button("🚪 Sair do Sistema"):
        st.session_state.autenticado = False
        st.session_state.supabase.auth.sign_out()
        st.rerun()

prof = st.text_input("Nome do(a) Professor(a):", value="")
comp = st.text_input("Componente Curricular / Disciplina:")
turma = st.text_input("Ano / Turma:")
st.write("---")
arqs = st.file_uploader("Arquivos de Referência Pedagógica (Até 3):", type=["pdf", "docx", "txt"], accept_multiple_files=True)

txt_ref = ""
if arqs and len(arqs) <= 3:
    for arq in arqs:
        try:
            if arq.type == "text/plain": txt_ref += arq.read().decode("utf-8") + "\n\n"
            elif arq.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                txt_ref += "\n".join([p.text for p in Document(arq).paragraphs]) + "\n\n"
            elif arq.type == "application/pdf":
                for pg in PdfReader(arq).pages: txt_ref += (pg.extract_text() or "") + "\n"
            st.success(f"✅ {arq.name} carregado para planejamento!")
        except: st.error("Erro ao ler arquivo.")

# 🧠 FORMATADOR DO WORD
def aplicar_formatacao_inteligente(paragrafo, texto_com_tags):
    paragrafo.text = ""
    partes = re.split(r'(\*\*.*?\*\*)', texto_com_tags)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            texto_limpo = parte[2:-2]
            if texto_limpo:
                run = paragrafo.add_run(texto_limpo)
                run.bold = True
        elif parte:
            run = paragrafo.add_run(parte)
    try:
        for r in paragrafo.runs: r.font.name, r.font.size, r.font.color.rgb = 'Arial', Pt(12), RGBColor(0,0,0)
    except: pass

def preencher_word(nome_modelo, dados_tags):
    if not os.path.exists(nome_modelo): return None
    doc = Document(nome_modelo)
    def processar_paragrafo(p):
        for tg, tx in dados_tags.items():
            if tg in p.text:
                if tg in ["{{CORPO_PROVA}}", "{{TEXTO_RELATORIO}}"] or tg.startswith("{{"):
                    texto_final = p.text.replace(tg, tx)
                    aplicar_formatacao_inteligente(p, texto_final)
                else:
                    p.text = p.text.replace(tg, tx)
                    for r in p.runs: r.font.name, r.font.size, r.font.color.rgb = 'Arial', Pt(12), RGBColor(0,0,0)
    for p in doc.paragraphs: processar_paragrafo(p)
    for t in doc.tables:
        for l in t.rows:
            for c in l.cells:
                for p in c.paragraphs: processar_paragrafo(p)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

st.write("---")
aba1, aba2, aba3, aba4, aba_chat, aba_cal = st.tabs(["📅 Plano Anual", "📝 Plano Mensal/Quinzenal", "✍️ Avaliações/Atividades", "📊 Relatórios", "💬 Conversar com a B.IA", "📆 Calendário da Escola"])
cfg = types.GenerateContentConfig()

with aba1:
    if st.button("✨ GERAR PLANO ANUAL", key="b1"):
        with st.spinner("Processando..."):
            pt = f"Crie um plano anual de {comp} ({turma}). Use as tags: [COMPETENCIAS_GERAIS], [COMPETENCIAS_ESPECIFICAS], [CONCEITOS1], [OBJETO1], [HABILIDADES1], [CONCEITOS2], [OBJETO2], [HABILIDADES2], [CONCEITOS3], [OBJETO3], [HABILIDADES3], [INSTRUMENTOS], [REFERENCIAS].\n\nREF:\n{txt_ref}"
            st.session_state.resultado = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=pt, config=cfg).text
            st.session_state.modelo_atual = "modelo_anual.docx"

with aba2:
    mes = st.selectbox("Selecione o Mês:", ["Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"])
    duracao = st.text_input("Tempo de duração:", value="15 dias")
    cmd_mensal = st.text_input("Foco temático do plano:")
    if st.button("✨ GERAR ESTRUTURA MENSAL", key="b2"):
        with st.spinner("Processando..."):
            pt = f"Crie um plano de aula para {mes} de {comp} ({turma}). Foco: {cmd_mensal}. Separe por: [AREA], [HABILIDADES], [OBJETO], [CRITERIOS], [METODOLOGIA], [INSTRUMENTOS], [REFERENCIAS].\n\nREF:\n{txt_ref}"
            st.session_state.resultado = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=pt, config=cfg).text
            st.session_state.modelo_atual = "modelo_mensal.docx"
            st.session_state.duracao_input = duracao

with aba3:
    col1, col2 = st.columns(2)
    with col1:
        t_av = st.selectbox("Tipo de Atividade:", ["Prova Objetiva/Discursiva", "Recuperação Paralela", "Trabalho Dirigido", "Lista de Exercícios"])
        qst = st.slider("Quantidade de Questões:", 1, 15, 5)
    with col2:
        nivel = st.selectbox("Nível de Dificuldade:", ["Fácil", "Médio", "Difícil", "Personalizado"])
        detalhes_personalizados = st.text_input("Especificações do nível:") if nivel == "Personalizado" else ""

    if st.button("✨ GERAR ATIVIDADE COM BASE NO GEMINI", key="b3"):
        with st.spinner("A B.IA está formulando suas questões..."):
            dif_texto = detalhes_personalizados if nivel == "Personalizado" else nivel
            pt = f"Crie uma atividade do tipo {t_av} com {qst} questões de {comp} ({turma}). Nível: {dif_texto}. Inclua GABARITO detalhado. Marque negritos com asteriscos duplos."
            if txt_ref: pt += f"\n\nUse como base:\n{txt_ref}"
            st.session_state.resultado = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=pt, config=cfg).text
            st.session_state.modelo_atual = "modelo_avaliacao.docx"

with aba4:
    t_re = st.selectbox("Tipo Relatório:", ["Desempenho da Turma", "Aluno PDI/AEE"])
    ctx = st.text_area("Contexto do Aluno/Turma:")
    if st.button("✨ GERAR RELATÓRIO", key="b4"):
        with st.spinner("Processando..."):
            pt = f"Escreva um relatório do tipo {t_re} para {comp} ({turma}). Contexto: {ctx}."
            st.session_state.resultado = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=pt, config=cfg).text
            st.session_state.modelo_atual = "modelo_avaliacao.docx"

with aba_chat:
    st.subheader("💬 Sala de Conversa com a B.IA")
    for msg in st.session_state.historico_chat:
        with st.chat_message("user" if msg["role"] == "user" else "assistant"): st.write(msg["text"])
    if prompt := st.chat_input("Digite sua mensagem para a B.IA..."):
        with st.chat_message("user"): st.write(prompt)
        st.session_state.historico_chat.append({"role": "user", "text": prompt})
        with st.spinner("B.IA está pensando..."):
            conversa_formatada = [types.Content(role=m["role"], parts=[types.Part.from_text(text=m["text"])]) for m in st.session_state.historico_chat]
            resposta_gemini = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=conversa_formatada).text
            with st.chat_message("assistant"): st.write(resposta_gemini)
            st.session_state.historico_chat.append({"role": "model", "text": resposta_gemini})
            st.rerun()

# --- 📆 PAINEL DE HORÁRIOS ULTRA-ESTÁVEL (EXTRAÇÃO VIA PYTHON) ---
with aba_cal:
    st.subheader("📆 Painel de Horários Oficial (EEB Pref Olegário Bernardes)")
    caminho_pdf = "calendario_eeb_olegario.pdf"
    
    if os.path.exists(caminho_pdf):
        # O Python extrai o texto localmente na máquina, evitando erros de API do Google
        if "texto_calendario_extraido" not in st.session_state:
            try:
                leitor = PdfReader(caminho_pdf)
                conteudo_completo = ""
                for pagina in leitor.pages:
                    conteudo_completo += (pagina.extract_text() or "") + "\n"
                st.session_state.texto_calendario_extraido = conteudo_completo
            except:
                st.session_state.texto_calendario_extraido = ""

        if st.session_state.texto_calendario_extraido:
            st.info("✅ Base de dados do Urânia carregada com sucesso diretamente do sistema escolar!")
            st.write("🔍 **Consulte a grade semanal por Professor ou por Turma/Aluno:**")
            termo_busca = st.text_input("Digite o nome do Professor ou o número da Turma (Ex: BIANCA, ADRIANO, VIVI, 301, 108):").upper()
            
            if termo_busca:
                with st.spinner(f"B.IA localizando os horários para '{termo_busca}'..."):
                    prompt_filtro = (
                        f"Com base na seguinte matriz de dados textuais extraída do software Urânia, "
                        f"filtre as informações e monte uma tabela clássica de horários semanais (Linhas: 1º a 6º aula | Colunas: Segunda a Sexta) "
                        f"mostrando onde o termo '{termo_busca}' aparece. "
                        f"Se for um PROFESSOR, mostre suas turmas. Se for uma TURMA, mostre as disciplinas e os nomes dos professores de cada aula.\n\n"
                        f"DADOS DA GRADE:\n{st.session_state.texto_calendario_extraido}"
                    )
                    
                    resposta_filtro = st.session_state.client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=prompt_filtro
                    ).text
                    st.markdown(resposta_filtro)
        else:
            st.error("Erro técnico ao ler o conteúdo do arquivo PDF local.")
    else:
        st.warning("⚠️ O arquivo 'calendario_eeb_olegario.pdf' não foi detectado no seu repositório. Por favor, suba o PDF no GitHub para ativar.")

# --- PAINEL DE VISUALIZAÇÃO DE DOCUMENTOS GERAIS ---
if st.session_state.resultado:
    st.write("---")
    st.subheader("🖥️ Tela da B.IA: Base de Dados Gemini em Tempo Real")
    texto_editado = st.text_area("Você pode revisar ou ajustar o texto abaixo diretamente:", value=st.session_state.resultado, height=350)
    tags_map = {"{{PROFESSOR}}": prof, "{{COMPONENTE}}": comp, "{{TURMA}}": turma, "{{CORPO_PROVA}}": texto_editado, "{{TEXTO_RELATORIO}}": texto_editado}
    
    if st.session_state.modelo_atual == "modelo_anual.docx":
        def ext(tg, tx):
            m = re.search(rf"\[{tg}\](.*?)(?=\[\w+\]|$)", tx, re.DOTALL)
            return m.group(1).strip() if m else "Em branco."
        tags_map.update({"{{COMPETENCIAS_GERAIS}}": ext("COMPETENCIAS_GERAIS", texto_editado), "{{COMPETENCIAS_ESPECIFICAS}}": ext("COMPETENCIAS_ESPECIFICAS", texto_editado), "{{CONCEITOS1}}": ext("CONCEITOS1", texto_editado), "{{OBJETO_CONHECIMENTO1}}": ext("OBJETO1", texto_editado), "{{HABILIDADES1}}": ext("HABILIDADES1", texto_editado), "{{CONCEITOS2}}": ext("CONCEITOS2", texto_editado), "{{OBJETO_CONHECIMENTO2}}": ext("OBJETO2", texto_editado), "{{HABILIDADES2}}": ext("HABILIDADES2", texto_editado), "{{CONCEITOS3}}": ext("CONCEITOS3", texto_editado), "{{OBJETO_CONHECIMENTO3}}": ext("OBJETO3", texto_editado), "{{HABILIDADES3}}": ext("HABILIDADES3", texto_editado), "{{INSTRUMENTOS}}": ext("INSTRUMENTOS", texto_editado), "{{REFERENCIAS}}": ext("REFERENCIAS", texto_editado)})
    elif st.session_state.modelo_atual == "modelo_mensal.docx":
        def ext_m(tg, tx):
            m = re.search(rf"\[{tg}\](.*?)(?=\[\w+\]|$)", tx, re.DOTALL)
            return m.group(1).strip() if m else "Em branco."
        tags_map.update({"{{AREA_CONHECIMENTO}}": ext_m("AREA", texto_editado), "{{HABILIDADES_MENSAL}}": ext_m("HABILIDADES", texto_editado), "{{OBJETO_MENSAL}}": ext_m("OBJETO", texto_editado), "{{CRITERIOS_MENSAL}}": ext_m("CRITERIOS", texto_editado), "{{METODOLOGIA_MENSAL}}": ext_m("METODOLOGIA", texto_editado), "{{INSTRUMENTOS_MENSAL}}": ext_m("INSTRUMENTOS", texto_editado), "{{DURACAO_MENSAL}}": st.session_state.get("duracao_input", "15 dias"), "{{REFERENCIAS_MENSAL}}": ext_m("REFERENCIAS", texto_editado)})
    
    w_bytes = preencher_word(st.session_state.modelo_atual, tags_map)
    if w_bytes: st.download_button("📥 BAIXAR DOCUMENTO NO MODELO OFICIAL (.DOCX)", data=w_bytes, file_name=f"Documento_{comp}.docx", key="dl_f")
from docx.shared import Pt, RGBColor
from google import genai
from google.genai import types
from pypdf import PdfReader
from supabase import create_client, Client
import io, os, re

st.set_page_config(page_title="Gerador Olegário Pro", layout="wide")

SUPABASE_URL = "https://hgdpahtyvoridomkktkt.supabase.co"
SUPABASE_KEY = "sb_publishable_yxUHijFIrK7h2yznh8ySKw_4eBFxAlI"
MINHA_CHAVE = "AQ.Ab8RN6LW-l2a1Qvu-8ZCD75G0OlT5RSVhxJpS3Pa6VcNQ88AFg"

if "supabase" not in st.session_state: st.session_state.supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
if "autenticado" not in st.session_state: st.session_state.autenticado = False
if "client" not in st.session_state: st.session_state.client = genai.Client(api_key=MINHA_CHAVE)
if "resultado" not in st.session_state: st.session_state.resultado = ""
if "historico_chat" not in st.session_state:
    st.session_state.historico_chat = [
        {"role": "model", "text": "Olá! Eu sou a B.IA, sua assistente pedagógica e companheira de trabalho. Sobre o que você gostaria de conversar ou planejar hoje? Estou pronta para ajudar em qualquer assunto!"}
    ]

# --- TELA DE LOGIN ---
if not st.session_state.autenticado:
    st.title("🔐 B.IA: Assistente Pedagógica")
    aba_login, aba_cadastro = st.tabs(["Entrar no Sistema", "Cadastrar Novo Professor"])
    with aba_login:
        email = st.text_input("E-mail funcional:", key="login_email")
        senha = st.text_input("Senha de acesso:", type="password", key="login_senha")
        if st.button("🔓 VALIDAR CREDENCIAIS", key="btn_login_exec"):
            try:
                res = st.session_state.supabase.auth.sign_in_with_password({"email": email, "password": senha})
                if res.session:
                    st.session_state.autenticado = True
                    st.success("✅ Credenciais validadas com sucesso!")
            except:
                st.error("E-mail ou senha incorretos no banco de dados.")
        if st.session_state.autenticado:
            if st.button("🚀 CLIQUE AQUI PARA ENTRAR NO PAINEL"):
                st.rerun()
    with aba_cadastro:
        novo_email = st.text_input("E-mail do Professor:", key="cad_email")
        nova_senha = st.text_input("Crie uma Senha (mínimo 6 dígitos):", type="password", key="cad_senha")
        if st.button("📝 CADASTRAR PROFESSOR", key="btn_cad_exec"):
            try:
                st.session_state.supabase.auth.sign_up({"email": novo_email, "password": nova_senha})
                st.success("Professor cadastrado com sucesso! Volte à aba anterior.")
            except Exception as e: st.error(f"Erro ao cadastrar: {e}")
    st.stop()

# --- INTERFACE PRINCIPAL ---
st.title("🧠 B.IA: Assistente Pedagógica")
with st.sidebar:
    st.header("⚙️ Painel do Professor")
    net = st.toggle("🔓 Liberar Internet (Google Search)", value=False)
    st.write("---")
    if st.button("🚪 Sair do Sistema"):
        st.session_state.autenticado = False
        st.session_state.supabase.auth.sign_out()
        st.rerun()

prof = st.text_input("Nome do(a) Professor(a):", value="")
comp = st.text_input("Componente Curricular / Disciplina:")
turma = st.text_input("Ano / Turma:")
st.write("---")
arqs = st.file_uploader("Arquivos de Referência Pedagógica (Até 3):", type=["pdf", "docx", "txt"], accept_multiple_files=True)

txt_ref = ""
if arqs and len(arqs) <= 3:
    for arq in arqs:
        try:
            if arq.type == "text/plain": txt_ref += arq.read().decode("utf-8") + "\n\n"
            elif arq.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                txt_ref += "\n".join([p.text for p in Document(arq).paragraphs]) + "\n\n"
            elif arq.type == "application/pdf":
                for pg in PdfReader(arq).pages: txt_ref += (pg.extract_text() or "") + "\n"
            st.success(f"✅ {arq.name} carregado para planejamento!")
        except: st.error("Erro ao ler arquivo.")

# 🧠 FORMATADOR DO WORD
def aplicar_formatacao_inteligente(paragrafo, texto_com_tags):
    paragrafo.text = ""
    partes = re.split(r'(\*\*.*?\*\*)', texto_com_tags)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            texto_limpo = parte[2:-2]
            if texto_limpo:
                run = paragrafo.add_run(texto_limpo)
                run.bold = True
        elif parte:
            run = paragrafo.add_run(parte)
    try:
        for r in paragrafo.runs: r.font.name, r.font.size, r.font.color.rgb = 'Arial', Pt(12), RGBColor(0,0,0)
    except: pass

def preencher_word(nome_modelo, dados_tags):
    if not os.path.exists(nome_modelo): return None
    doc = Document(nome_modelo)
    def processar_paragrafo(p):
        for tg, tx in dados_tags.items():
            if tg in p.text:
                if tg in ["{{CORPO_PROVA}}", "{{TEXTO_RELATORIO}}"] or tg.startswith("{{"):
                    texto_final = p.text.replace(tg, tx)
                    aplicar_formatacao_inteligente(p, texto_final)
                else:
                    p.text = p.text.replace(tg, tx)
                    for r in p.runs: r.font.name, r.font.size, r.font.color.rgb = 'Arial', Pt(12), RGBColor(0,0,0)
    for p in doc.paragraphs: processar_paragrafo(p)
    for t in doc.tables:
        for l in t.rows:
            for c in l.cells:
                for p in c.paragraphs: processar_paragrafo(p)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

st.write("---")
aba1, aba2, aba3, aba4, aba_chat, aba_cal = st.tabs(["📅 Plano Anual", "📝 Plano Mensal/Quinzenal", "✍️ Avaliações/Atividades", "📊 Relatórios", "💬 Conversar com a B.IA", "📆 Calendário da Escola"])
cfg = types.GenerateContentConfig()

with aba1:
    if st.button("✨ GERAR PLANO ANUAL", key="b1"):
        with st.spinner("Processando..."):
            pt = f"Crie um plano anual de {comp} ({turma}). Use as tags: [COMPETENCIAS_GERAIS], [COMPETENCIAS_ESPECIFICAS], [CONCEITOS1], [OBJETO1], [HABILIDADES1], [CONCEITOS2], [OBJETO2], [HABILIDADES2], [CONCEITOS3], [OBJETO3], [HABILIDADES3], [INSTRUMENTOS], [REFERENCIAS].\n\nREF:\n{txt_ref}"
            st.session_state.resultado = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=pt, config=cfg).text
            st.session_state.modelo_atual = "modelo_anual.docx"

with aba2:
    mes = st.selectbox("Selecione o Mês:", ["Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"])
    duracao = st.text_input("Tempo de duração:", value="15 dias")
    cmd_mensal = st.text_input("Foco temático do plano:")
    if st.button("✨ GERAR ESTRUTURA MENSAL", key="b2"):
        with st.spinner("Processando..."):
            pt = f"Crie um plano de aula para {mes} de {comp} ({turma}). Foco: {cmd_mensal}. Separe por: [AREA], [HABILIDADES], [OBJETO], [CRITERIOS], [METODOLOGIA], [INSTRUMENTOS], [REFERENCIAS].\n\nREF:\n{txt_ref}"
            st.session_state.resultado = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=pt, config=cfg).text
            st.session_state.modelo_atual = "modelo_mensal.docx"
            st.session_state.duracao_input = duracao

with aba3:
    col1, col2 = st.columns(2)
    with col1:
        t_av = st.selectbox("Tipo de Atividade:", ["Prova Objetiva/Discursiva", "Recuperação Paralela", "Trabalho Dirigido", "Lista de Exercícios"])
        qst = st.slider("Quantidade de Questões:", 1, 15, 5)
    with col2:
        nivel = st.selectbox("Nível de Dificuldade:", ["Fácil", "Médio", "Difícil", "Personalizado"])
        detalhes_personalizados = st.text_input("Especificações do nível:") if nivel == "Personalizado" else ""

    if st.button("✨ GERAR ATIVIDADE COM BASE NO GEMINI", key="b3"):
        with st.spinner("A B.IA está formulando suas questões..."):
            dif_texto = detalhes_personalizados if nivel == "Personalizado" else nivel
            pt = f"Crie uma atividade do tipo {t_av} com {qst} questões de {comp} ({turma}). Nível: {dif_texto}. Inclua GABARITO detalhado. Marque negritos com asteriscos duplos."
            if txt_ref: pt += f"\n\nUse como base:\n{txt_ref}"
            st.session_state.resultado = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=pt, config=cfg).text
            st.session_state.modelo_atual = "modelo_avaliacao.docx"

with aba4:
    t_re = st.selectbox("Tipo Relatório:", ["Desempenho da Turma", "Aluno PDI/AEE"])
    ctx = st.text_area("Contexto do Aluno/Turma:")
    if st.button("✨ GERAR RELATÓRIO", key="b4"):
        with st.spinner("Processando..."):
            pt = f"Escreva um relatório do tipo {t_re} para {comp} ({turma}). Contexto: {ctx}."
            st.session_state.resultado = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=pt, config=cfg).text
            st.session_state.modelo_atual = "modelo_avaliacao.docx"

with aba_chat:
    st.subheader("💬 Sala de Conversa com a B.IA")
    for msg in st.session_state.historico_chat:
        with st.chat_message("user" if msg["role"] == "user" else "assistant"): st.write(msg["text"])
    if prompt := st.chat_input("Digite sua mensagem para a B.IA..."):
        with st.chat_message("user"): st.write(prompt)
        st.session_state.historico_chat.append({"role": "user", "text": prompt})
        with st.spinner("B.IA está pensando..."):
            conversa_formatada = [types.Content(role=m["role"], parts=[types.Part.from_text(text=m["text"])]) for m in st.session_state.historico_chat]
            resposta_gemini = st.session_state.client.models.generate_content(model='gemini-2.5-flash', contents=conversa_formatada).text
            with st.chat_message("assistant"): st.write(resposta_gemini)
            st.session_state.historico_chat.append({"role": "model", "text": resposta_gemini})
            st.rerun()

# --- 📆 LEITURA DIRETA DO ARQUIVO PDF ESTÁTICO DO REPOSITÓRIO ---
with aba_cal:
    st.subheader("📆 Painel de Horários Oficial (EEB Pref Olegário Bernardes)")
    
    caminho_pdf = "calendario_eeb_olegario.pdf"
    
    if os.path.exists(caminho_pdf):
        st.info("✅ Base de dados do Urânia carregada com sucesso diretamente do sistema escolar!")
        st.write("🔍 **Consulte a grade semanal por Professor ou por Turma/Aluno:**")
        termo_busca = st.text_input("Digite o nome do Professor ou o número da Turma (Ex: BIANCA, ADRIANO, VIVI, 301, 108):").upper()
        
        if termo_busca:
            with st.spinner(f"B.IA localizando os horários para '{termo_busca}'..."):
                try:
                    with open(caminho_pdf, "rb") as f:
                        pdf_bytes = f.read()
                    
                    documento_inline = types.Part.from_bytes(data=pdf_bytes, mime_type="application/pdf")
                    
                    prompt_filtro = (
                        f"Analise o relatório completo de horários escolares do Urânia contido neste PDF. "
                        f"Filtre os dados com precisão absoluta e monte uma tabela clássica de horários semanais (com linhas de 1º a 6º aula "
                        f"e colunas de Segunda a Sexta) mostrando onde o termo '{termo_busca}' aparece. "
                        f"Se o termo for um PROFESSOR, liste as turmas dele. Se for uma TURMA, mostre as disciplinas e os nomes dos professores de cada aula."
                    )
                    
                    resposta_filtro = st.session_state.client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=[documento_inline, prompt_filtro]
                    ).text
                    st.markdown(resposta_filtro)
                except Exception as e:
                    st.error("Erro no processamento da inteligência artificial ao ler o arquivo.")
    else:
        st.warning("⚠️ O arquivo 'calendario_eeb_olegario.pdf' não foi encontrado na pasta do GitHub. Por favor, faça o upload dele para ativar este painel.")

# --- PAINEL DE VISUALIZAÇÃO DE DOCUMENTOS GERAIS ---
if st.session_state.resultado:
    st.write("---")
    st.subheader("🖥️ Tela da B.IA: Base de Dados Gemini em Tempo Real")
    texto_editado = st.text_area("Você pode revisar ou ajustar o texto abaixo diretamente:", value=st.session_state.resultado, height=350)
    tags_map = {"{{PROFESSOR}}": prof, "{{COMPONENTE}}": comp, "{{TURMA}}": turma, "{{CORPO_PROVA}}": texto_editado, "{{TEXTO_RELATORIO}}": texto_editado}
    
    if st.session_state.modelo_atual == "modelo_anual.docx":
        def ext(tg, tx):
            m = re.search(rf"\[{tg}\](.*?)(?=\[\w+\]|$)", tx, re.DOTALL)
            return m.group(1).strip() if m else "Em branco."
        tags_map.update({"{{COMPETENCIAS_GERAIS}}": ext("COMPETENCIAS_GERAIS", texto_editado), "{{COMPETENCIAS_ESPECIFICAS}}": ext("COMPETENCIAS_ESPECIFICAS", texto_editado), "{{CONCEITOS1}}": ext("CONCEITOS1", texto_editado), "{{OBJETO_CONHECIMENTO1}}": ext("OBJETO1", texto_editado), "{{HABILIDADES1}}": ext("HABILIDADES1", texto_editado), "{{CONCEITOS2}}": ext("CONCEITOS2", texto_editado), "{{OBJETO_CONHECIMENTO2}}": ext("OBJETO2", texto_editado), "{{HABILIDADES2}}": ext("HABILIDADES2", texto_editado), "{{CONCEITOS3}}": ext("CONCEITOS3", texto_editado), "{{OBJETO_CONHECIMENTO3}}": ext("OBJETO3", texto_editado), "{{HABILIDADES3}}": ext("HABILIDADES3", texto_editado), "{{INSTRUMENTOS}}": ext("INSTRUMENTOS", texto_editado), "{{REFERENCIAS}}": ext("REFERENCIAS", texto_editado)})
    elif st.session_state.modelo_atual == "modelo_mensal.docx":
        def ext_m(tg, tx):
            m = re.search(rf"\[{tg}\](.*?)(?=\[\w+\]|$)", tx, re.DOTALL)
            return m.group(1).strip() if m else "Em branco."
        tags_map.update({"{{AREA_CONHECIMENTO}}": ext_m("AREA", texto_editado), "{{HABILIDADES_MENSAL}}": ext_m("HABILIDADES", texto_editado), "{{OBJETO_MENSAL}}": ext_m("OBJETO", texto_editado), "{{CRITERIOS_MENSAL}}": ext_m("CRITERIOS", texto_editado), "{{METODOLOGIA_MENSAL}}": ext_m("METODOLOGIA", texto_editado), "{{INSTRUMENTOS_MENSAL}}": ext_m("INSTRUMENTOS", texto_editado), "{{DURACAO_MENSAL}}": st.session_state.get("duracao_input", "15 dias"), "{{REFERENCIAS_MENSAL}}": ext_m("REFERENCIAS", texto_editado)})
    
    w_bytes = preencher_word(st.session_state.modelo_atual, tags_map)
    if w_bytes: st.download_button("📥 BAIXAR DOCUMENTO NO MODELO OFICIAL (.DOCX)", data=w_bytes, file_name=f"Documento_{comp}.docx", key="dl_f")
